"""
FastAPI server for Universal Scraper - GCP Cloud Run deployment
"""
import os
import logging
import io
import time
import hashlib
from typing import List, Optional, Dict, Any
from fastapi import FastAPI, HTTPException, Header, BackgroundTasks, UploadFile, File, Form, Response, Query
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import uvicorn

from universal_scraper.core.scraper import UniversalScraper
from universal_scraper.core.direct_llm_extractor import DirectLLMExtractor
from universal_scraper.core.redis_cache import RedisCache
from universal_scraper.core.tenant_cache import TenantCache
from universal_scraper.core.tenant_pattern_cache import TenantPatternCache, CacheVisibility
from universal_scraper.core.agent_manager import AgentManager, AgentType, AgentStatus
from universal_scraper.core.field_discovery import FieldDiscovery
from api.middleware.auth import get_tenant_id, get_tenant_context, get_current_user
from api.middleware.rate_limit import RateLimiter
from api.middleware.usage_tracking import UsageTracker
from fastapi import Depends

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="Universal Scraper API",
    description="AI-powered universal web scraper API",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "https://universal-scaper.web.app",
        "https://universal-scaper.firebaseapp.com",
    ],
    allow_origin_regex=r"https?://(localhost|127\.0\.0\.1)(:\d+)?|https://.*\.web\.app|https://.*\.firebaseapp\.com",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Request models
class ScrapeRequest(BaseModel):
    url: str = Field(..., description="URL to scrape")
    fields: List[str] = Field(default=[], description="Fields to extract (empty = auto-extract)")
    target: Optional[str] = Field(default=None, description="What to look for (e.g., 'products', 'comments')")
    html: Optional[str] = Field(default=None, description="Optional pre-fetched HTML content")
    mode: str = Field(default="hybrid", description="Fetch mode: html, browser, hybrid")
    force_html: bool = Field(default=False, description="Skip JSON detection")
    force_generate: bool = Field(default=False, description="Skip cache, generate new code")
    scroll_to_bottom: bool = Field(default=False, description="Scroll to bottom for infinite scroll")
    click_load_more: Optional[str] = Field(default=None, description="CSS selector for Load More button")
    wait_for_selector: Optional[str] = Field(default=None, description="Wait for selector before scraping")
    proxy_config: Optional[Dict[str, Any]] = Field(default=None, description="Proxy configuration")
    browser_timeout: Optional[int] = Field(default=120000, description="Browser navigation timeout in milliseconds (default: 120s)")

class CrawlRequest(BaseModel):
    start_urls: List[str] = Field(..., description="Starting URLs for crawling")
    fields: List[str] = Field(default=[], description="Fields to extract")
    max_pages: int = Field(default=10, description="Maximum pages to crawl")
    max_depth: int = Field(default=2, description="Maximum crawl depth")
    follow_patterns: Optional[List[str]] = Field(default=None, description="URL patterns to follow")
    ignore_patterns: Optional[List[str]] = Field(default=None, description="URL patterns to ignore")

class HealthResponse(BaseModel):
    status: str
    version: str

class DocumentProcessingResponse(BaseModel):
    success: bool
    data: List[Dict[str, Any]]
    metadata: Dict[str, Any]

class ProxyTestRequest(BaseModel):
    provider: str
    server: Optional[str] = None
    username: Optional[str] = None
    password: Optional[str] = None
    country: Optional[str] = None
    apifyProxy: Optional[Dict[str, Any]] = None
    externalProxy: Optional[Dict[str, Any]] = None
    webUnblocker: Optional[Dict[str, Any]] = None  # For Web Unblocker testing

class ProxyTestResponse(BaseModel):
    success: bool
    message: str
    ip: Optional[str] = None

class SuggestFieldsRequest(BaseModel):
    url: str = Field(..., description="URL to analyze for field suggestions")
    target: Optional[str] = Field(default=None, description="What to look for (e.g., 'products', 'comments')")
    use_llm: bool = Field(default=False, description="Use LLM for field discovery (slower but more accurate)")
    proxy_config: Optional[Dict[str, Any]] = Field(default=None, description="Proxy configuration")
    browser_timeout: Optional[int] = Field(default=60000, description="Browser timeout in milliseconds")

class SuggestFieldsResponse(BaseModel):
    fields: List[str] = Field(..., description="Suggested field names")
    confidence: float = Field(..., description="Confidence score (0-1)")
    source: str = Field(..., description="Source of field discovery (json/html/llm)")
    reasoning: str = Field(..., description="Explanation of how fields were discovered")
    unblocker_log: Optional[List[Dict[str, Any]]] = Field(default=None, description="Log of unblocking steps")

class SavePatternRequest(BaseModel):
    url: str = Field(..., description="URL the pattern is for")
    fields: List[str] = Field(..., description="Fields extracted")
    pattern_data: Dict[str, Any] = Field(..., description="Extraction pattern (selectors or Direct LLM result)")
    visibility: str = Field(default="private", description="Pattern visibility: 'private' or 'public'")

class DeletePatternRequest(BaseModel):
    domain: str = Field(..., description="Domain of the pattern")
    fields: List[str] = Field(..., description="Fields of the pattern")

class GenerateCodeRequest(BaseModel):
    url: str = Field(..., description="URL of the page")
    fields: List[str] = Field(..., description="Fields to extract")
    selectors: Dict[str, Any] = Field(..., description="Deterministic selectors")
    target: Optional[str] = Field(default=None, description="Extraction target")

# Scraper pool (reuse instances)
scraper_pool: Dict[str, UniversalScraper] = {}
extractor_pool: Dict[str, DirectLLMExtractor] = {}


def convert_proxy_config(frontend_config: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """
    Convert frontend proxy configuration format to backend format.
    Handles nested frontend format, comma-separated strings, and Web Unblocker specific logic.
    """
    if not frontend_config:
        return None
        
    provider = frontend_config.get('provider', 'none')
    if provider == 'none':
        return None
        
    # Initialize result with defaults
    result = {
        'server': None,
        'username': None,
        'password': None,
        'web_unlocker_api_key': None,
        'web_unlocker_zone': 'web_unlocker1',
        'web_unlocker': False
    }

    # 1. Extract raw values from frontend config
    external_proxy = frontend_config.get('externalProxy', {})
    server = external_proxy.get('server')
    username = external_proxy.get('username')
    password = external_proxy.get('password')
    
    # 2. Robust parsing for comma-separated credentials (host,port,user,pass)
    # This handles the case where the user pastes the entire string into the 'server' field
    if server and ',' in server:
        parts = [p.strip() for p in server.split(',')]
        if len(parts) >= 4:
            server = f"{parts[0]}:{parts[1]}"
            username = parts[2]
            password = parts[3]
            logger.info(f"✅ Parsed comma-separated proxy: {server} (user: {username})")
        elif len(parts) == 2:
            server = f"{parts[0]}:{parts[1]}"
    
    # 3. Normalize server string
    if server:
        server = server.replace(',', ':')
        if not server.startswith('http'):
            server = f"http://{server}"
            
    # 4. Handle Web Unblocker specific logic
    if provider in ['web_unlocker', 'web_unblocker']:
        result['web_unlocker'] = True
        web_unblocker_config = frontend_config.get('webUnblocker', {})
        result['web_unlocker_zone'] = web_unblocker_config.get('zone', 'web_unlocker1')
        
        if web_unblocker_config.get('enabled'):
            if web_unblocker_config.get('useProxyMethod'):
                # Use the parsed credentials
                if server and username and password:
                    # Construct internal API key format: host:port:user:pass
                    host_port = server.replace('http://', '').replace('https://', '')
                    result['web_unlocker_api_key'] = f"{host_port}:{username}:{password}"
            else:
                # Use explicit API key
                result['web_unlocker_api_key'] = web_unblocker_config.get('apiKey')
                
    # 5. Default server for Bright Data if missing
    if not server and username and 'brd-customer' in username:
        server = 'http://brd.superproxy.io:33335'
        logger.info(f"✅ Defaulting to Bright Data server: {server}")
        
    # 6. Fallback for Web Unblocker if only API key provided via environment
    if result['web_unlocker'] and not server and not username:
        customer_id = os.getenv('WEB_UNBLOCKER_CUSTOMER_ID', 'hl_803e8195')
        server = 'http://brd.superproxy.io:33335'
        username = f'brd-customer-{customer_id}-zone-{result["web_unlocker_zone"]}'
        password = os.getenv('WEB_UNBLOCKER_API_KEY')
        logger.info(f"✅ Constructing Web Unblocker proxy from environment: {server}")

    # Set final values
    result['server'] = server
    result['username'] = username
    result['password'] = password
    
    # Return None if no server was found
    if not result['server']:
        return None
        
    return result

def get_extractor(api_key: str, redis_cache: Optional[RedisCache] = None) -> DirectLLMExtractor:
    """Get or create extractor instance with Redis cache support"""
    cache_key = f"{api_key[:10]}"
    if cache_key not in extractor_pool:
        extractor_pool[cache_key] = DirectLLMExtractor(
            api_key=api_key,
            model_name="gpt-4o-mini",
            quality_mode="balanced",
            redis_cache=redis_cache or get_redis_cache()  # Use Redis cache for multi-tenant SaaS
        )
    return extractor_pool[cache_key]

async def extract_text_from_file(file: UploadFile, use_ocr: bool = False) -> str:
    """Extract text from uploaded file"""
    content = await file.read()
    filename = file.filename.lower()
    
    # PDF
    if filename.endswith('.pdf'):
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            if use_ocr:
                # TODO: Implement OCR for PDF
                raise HTTPException(status_code=400, detail=f"PDF text extraction failed: {str(e)}")
            raise HTTPException(status_code=400, detail=f"PDF text extraction failed: {str(e)}")
    
    # DOCX
    elif filename.endswith(('.doc', '.docx')):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join([cell.text for cell in row.cells])
            return text
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise HTTPException(status_code=400, detail=f"DOCX text extraction failed: {str(e)}")
    
    # TXT, MD
    elif filename.endswith(('.txt', '.md')):
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return content.decode('latin-1')
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Text file decoding failed: {str(e)}")
    
    # Images (OCR)
    elif filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')) and use_ocr:
        try:
            import pytesseract
            from PIL import Image
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            raise HTTPException(status_code=400, detail=f"OCR extraction failed: {str(e)}")
    
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}")

def get_scraper(
    api_key: str,
    mode: str = "hybrid",
    proxy_config: Optional[Dict[str, Any]] = None,
    redis_cache: Optional[RedisCache] = None,
    browser_timeout: int = 120000  # Default 120 seconds for slow-loading pages
) -> UniversalScraper:
    """
    Get or create scraper instance
    """
    # 1. Convert proxy config to backend format (Robust parsing)
    backend_proxy_config = convert_proxy_config(proxy_config)
    
    # 2. Extract web unlocker details from the converted config
    web_unblocker_api_key = None
    web_unblocker_zone = "web_unlocker1"
    
    if backend_proxy_config:
        web_unblocker_api_key = backend_proxy_config.get('web_unlocker_api_key')
        web_unblocker_zone = backend_proxy_config.get('web_unlocker_zone', 'web_unlocker1')
    
    # Fallback to environment variables if not provided in request
    if not web_unblocker_api_key:
        web_unblocker_api_key = os.getenv("WEB_UNBLOCKER_API_KEY")
        web_unblocker_zone = os.getenv("WEB_UNBLOCKER_ZONE", "web_unlocker1")
        if web_unblocker_api_key:
            logger.info(f"🌐 Using Web Unblocker from environment (zone: {web_unblocker_zone})")
    
    # 3. Include proxy config and timeout in key for proper pooling
    proxy_key = ""
    if backend_proxy_config:
        proxy_key = f":{backend_proxy_config.get('server', '')[:20]}"
    web_unlocker_key = f":web_unlocker{web_unblocker_api_key[:10] if web_unblocker_api_key else 'none'}"
    timeout_key = f":timeout{browser_timeout}"
    key = f"{api_key[:10]}:{mode}{proxy_key}{web_unlocker_key}{timeout_key}"
    
    if key not in scraper_pool:
        scraper_pool[key] = UniversalScraper(
            api_key=api_key,
            fetch_mode=mode,
            enable_cache=True,
            use_camoufox=True,  # Enable Camoufox by default for superior anti-detection
            proxy_config=backend_proxy_config,  # Pass converted proxy config to scraper
            redis_cache=redis_cache,  # Pass Redis cache for multi-tenant SaaS
            browser_timeout=browser_timeout,  # Pass browser timeout
            web_unblocker_api_key=web_unblocker_api_key,  # Pass web unlocker API key
            web_unblocker_zone=web_unblocker_zone  # Pass web unlocker zone
        )
    return scraper_pool[key]

@app.get("/", response_model=HealthResponse)
async def root():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "version": "1.0.0"
    }

@app.get("/health", response_model=HealthResponse)
async def health():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "version": "1.0.0"
    }

# Initialize shared Redis cache and services
_redis_cache = None
_rate_limiter = None
_usage_tracker = None
_tenant_pattern_cache = None
_agent_manager = None

def get_redis_cache() -> RedisCache:
    """Get or create Redis cache instance"""
    global _redis_cache
    if _redis_cache is None:
        _redis_cache = RedisCache()
    return _redis_cache

def get_rate_limiter() -> RateLimiter:
    """Get or create rate limiter"""
    global _rate_limiter
    if _rate_limiter is None:
        _rate_limiter = RateLimiter(get_redis_cache())
    return _rate_limiter

def get_usage_tracker() -> UsageTracker:
    """Get or create usage tracker"""
    global _usage_tracker
    if _usage_tracker is None:
        _usage_tracker = UsageTracker(get_redis_cache())
    return _usage_tracker

def get_tenant_pattern_cache() -> TenantPatternCache:
    """Get or create tenant pattern cache"""
    global _tenant_pattern_cache
    if _tenant_pattern_cache is None:
        _tenant_pattern_cache = TenantPatternCache(get_redis_cache())
    return _tenant_pattern_cache

def get_agent_manager() -> AgentManager:
    """Get or create agent manager"""
    global _agent_manager
    if _agent_manager is None:
        cloud_tasks_enabled = os.getenv("CLOUD_TASKS_ENABLED", "false").lower() == "true"
        _agent_manager = AgentManager(get_redis_cache(), cloud_tasks_enabled=cloud_tasks_enabled)
    return _agent_manager

class GenerateFieldsFromPromptRequest(BaseModel):
    prompt: str = Field(..., description="Natural language description of what data to extract")
    url: Optional[str] = Field(default=None, description="Optional URL for domain context")

class GenerateFieldsFromPromptResponse(BaseModel):
    fields: List[str] = Field(..., description="Generated field names")
    descriptions: Optional[Dict[str, str]] = Field(default=None, description="Field descriptions if requested")

@app.post("/api/v1/generate-fields-from-prompt", response_model=GenerateFieldsFromPromptResponse)
async def generate_fields_from_prompt_endpoint(
    request: GenerateFieldsFromPromptRequest,
    tenant_id: str = Depends(get_tenant_id),
    current_user: dict = Depends(get_current_user),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
) -> GenerateFieldsFromPromptResponse:
    """
    Generate field names from a natural language prompt (no URL required).
    
    This endpoint uses LLM to convert natural language descriptions into structured field names.
    Useful for suggesting fields before navigating to a page.
    
    Example:
        Prompt: "I want product names, prices in USD, and star ratings"
        Returns: ["product_name", "price", "rating"]
    """
    # Get API key from multiple sources (priority order)
    api_key = None
    
    # 1. From X-API-Key header (highest priority - direct parameter)
    if x_api_key:
        api_key = x_api_key
        logger.debug(f"[{tenant_id}] Using API key from X-API-Key header")
    
    # 2. From current_user dict (if provided via dependency)
    if not api_key and current_user and current_user.get('api_key'):
        api_key = current_user.get('api_key')
        logger.debug(f"[{tenant_id}] Using API key from current_user")
    
    # 3. From environment variable (fallback)
    if not api_key:
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            logger.debug(f"[{tenant_id}] Using API key from environment variable")
    
    if not api_key:
        logger.warning(f"[{tenant_id}] No API key found in headers or environment")
        raise HTTPException(
            status_code=400, 
            detail="OpenAI API key required for field generation. Provide X-API-Key header or set OPENAI_API_KEY environment variable"
        )
    
    try:
        logger.info(f"[{tenant_id}] Generating fields from prompt: {request.prompt[:80]}...")
        
        # Use UniversalScraper's static method to generate fields
        fields = await UniversalScraper.generate_fields_from_prompt(
            prompt=request.prompt,
            url=request.url,
            api_key=api_key,
            return_descriptions=False
        )
        
        logger.info(f"[{tenant_id}] Generated {len(fields)} fields: {', '.join(fields)}")
        
        return GenerateFieldsFromPromptResponse(
            fields=fields,
            descriptions=None
        )
    except Exception as e:
        logger.error(f"[{tenant_id}] Field generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Field generation failed: {str(e)}")

@app.post("/api/v1/suggest-fields", response_model=SuggestFieldsResponse)
async def suggest_fields_endpoint(
    request: SuggestFieldsRequest,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Suggest fields to extract from a webpage (lightweight analysis, no full scrape)
    
    This endpoint analyzes the page structure to suggest relevant fields without
    performing a full scrape. Much faster than scraping.
    """
    # Get API key (optional for field discovery, but needed if use_llm=True)
    api_key = x_api_key or os.getenv("OPENAI_API_KEY") or os.getenv("GEMINI_API_KEY")
    
    if request.use_llm and not api_key:
        raise HTTPException(
            status_code=400,
            detail="API key required when use_llm=true. Provide X-API-Key header or set OPENAI_API_KEY environment variable"
        )
    
    try:
        logger.info(f"[{tenant_id}] Field discovery for URL: {request.url}")
        
        # 1. Convert frontend proxy config format to backend format (Robust parsing)
        backend_proxy_config = convert_proxy_config(request.proxy_config) if request.proxy_config else None
        
        # 2. Extract Web Unblocker details from the converted config
        web_unblocker_api_key = None
        web_unblocker_zone = "web_unlocker1"
        
        if backend_proxy_config:
            web_unblocker_api_key = backend_proxy_config.get('web_unlocker_api_key')
            web_unblocker_zone = backend_proxy_config.get('web_unlocker_zone', 'web_unlocker1')
        
        # Fallback to environment variables if not provided in request
        if not web_unblocker_api_key:
            web_unblocker_api_key = os.getenv("WEB_UNBLOCKER_API_KEY")
            web_unblocker_zone = os.getenv("WEB_UNBLOCKER_ZONE", "web_unlocker1")
            if web_unblocker_api_key:
                logger.info(f"🌐 Using Web Unblocker from environment (zone: {web_unblocker_zone})")
        
        # HybridFetcher now handles domain-specific forcing (e.g., Home Depot)
        from universal_scraper.core.hybrid_fetcher import HybridFetcher
        force_mode = None

        hybrid_fetcher = HybridFetcher(
            proxy_config=backend_proxy_config,
            headless=True,
            browser_timeout=request.browser_timeout or 90000, # Increased timeout for discovery
            force_mode=force_mode,
            use_camoufox=True,     # Enable Camoufox for superior anti-detection fallback
            web_unblocker_api_key=web_unblocker_api_key,  # Pass Web Unblocker for anti-bot
            web_unblocker_zone=web_unblocker_zone  # Pass Web Unblocker zone
        )
        
        try:
            fetch_result = await hybrid_fetcher.fetch(request.url)
            html = fetch_result.get('html', '')
        except Exception as e:
            logger.warning(f"Hybrid fetch failed, trying static: {e}")
            # Fallback to static HTML fetcher
            from universal_scraper.core.html_fetcher import HTMLFetcher
            try:
                html_fetcher = HTMLFetcher(proxy_config=backend_proxy_config)
                fetch_result = html_fetcher.fetch(request.url)
                html = fetch_result.get('html', '')
            except Exception as e2:
                logger.error(f"Static fetch also failed: {e2}")
                raise HTTPException(status_code=400, detail=f"Failed to fetch HTML content: {str(e2)}")
        
        if not html or len(html) < 100:
            raise HTTPException(status_code=400, detail="Failed to fetch HTML content or content too small")
        
        # Discover fields
        # Use LLM by default for better accuracy, unless explicitly disabled
        use_llm = request.use_llm if hasattr(request, 'use_llm') else (api_key is not None)
        field_discovery = FieldDiscovery(api_key=api_key if use_llm else None)
        result = await field_discovery.discover_fields(html, request.url, use_llm=use_llm, target=request.target)
        
        return SuggestFieldsResponse(
            fields=result['fields'],
            confidence=result['confidence'],
            source=result['source'],
            reasoning=result.get('reasoning', 'Field discovery completed'),
            unblocker_log=fetch_result.get('unblocker_log', [])
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[{tenant_id}] Field discovery failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Field discovery failed: {str(e)}")


# ============================================================================
# VISUAL PREVIEW ENDPOINT - Renders page with element highlighting
# ============================================================================

class PreviewRequest(BaseModel):
    url: str = Field(..., description="URL to preview")
    fields: Optional[List[str]] = Field(default=None, description="Fields to highlight")
    proxy_config: Optional[Dict[str, Any]] = Field(default=None, description="Proxy configuration")
    browser_timeout: Optional[int] = Field(default=60000, description="Browser timeout in milliseconds")

class DetectedElement(BaseModel):
    field: str = Field(..., description="Field name")
    selector: str = Field(..., description="CSS selector")
    sample_value: Optional[str] = Field(default=None, description="Sample extracted value")
    count: int = Field(default=1, description="Number of matching elements")

class PreviewResponse(BaseModel):
    success: bool
    html: str = Field(..., description="HTML with injected highlighting script")
    detected_elements: List[DetectedElement] = Field(default=[], description="Detected extractable elements")
    json_sources: Optional[List[Dict[str, Any]]] = Field(default=None, description="Discovered JSON sources")
    json_recommended: bool = Field(default=False, description="Whether JSON extraction is recommended")
    json_confidence: float = Field(default=0.0, description="Confidence score for JSON extraction (0-1)")
    extraction_mode: str = Field(default="browser", description="Recommended extraction mode: 'json' or 'browser'")
    fetch_method: str = Field(default="unknown", description="How the page was fetched: 'browser', 'static', 'static_fallback'")
    browser_rendering_failed: bool = Field(default=False, description="Whether browser rendering failed and static HTML was used")
    fallback_reason: Optional[str] = Field(default=None, description="Reason for fallback if browser failed")
    page_info: Dict[str, Any] = Field(default={}, description="Page metadata")
    unblocker_log: Optional[List[Dict[str, Any]]] = None

@app.post("/api/v1/preview", response_model=PreviewResponse)
async def preview_endpoint(
    request: PreviewRequest,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Get a preview of a webpage with element highlighting for visual field selection.
    
    Returns:
    - HTML with injected highlighting script
    - Detected extractable elements with CSS selectors
    - Discovered JSON sources
    - Page metadata
    """
    try:
        logger.info(f"[{tenant_id}] Preview request for: {request.url}")
        
        # 1. Convert frontend proxy config format to backend format (Robust parsing)
        backend_proxy_config = convert_proxy_config(request.proxy_config) if request.proxy_config else None
        
        # 2. Extract Web Unblocker details from the converted config
        web_unblocker_api_key = None
        web_unblocker_zone = "web_unlocker1"
        
        if backend_proxy_config:
            web_unblocker_api_key = backend_proxy_config.get('web_unlocker_api_key')
            web_unblocker_zone = backend_proxy_config.get('web_unlocker_zone', 'web_unlocker1')
        
        # Fallback to environment variables if not provided in request
        if not web_unblocker_api_key:
            web_unblocker_api_key = os.getenv("WEB_UNBLOCKER_API_KEY")
            web_unblocker_zone = os.getenv("WEB_UNBLOCKER_ZONE", "web_unlocker1")
            if web_unblocker_api_key:
                logger.info(f"🌐 Using Web Unblocker from environment (zone: {web_unblocker_zone})")
        
        proxy_config = backend_proxy_config
        
        # For preview, force browser mode to ensure full JS rendering
        # This ensures users see the actual rendered page, not just static HTML shell
        # CRITICAL: Pass Web Unblocker config to bypass Cloudflare/anti-bot protection
        # IMPORTANT: Use longer timeout and wait for JavaScript to fully render
        from universal_scraper.core.hybrid_fetcher import HybridFetcher
        from universal_scraper.core.json_detector import JSONDetector
        from bs4 import BeautifulSoup
        import json as json_module
        
        fetcher = HybridFetcher(
            proxy_config=proxy_config,
            headless=True,
            browser_timeout=max(request.browser_timeout or 60000, 90000),  # At least 90s for JS-heavy sites
            use_camoufox=True,  # Enable Camoufox for superior anti-detection
            force_mode='browser',  # Force browser mode for preview to show full rendered content
            web_unblocker_api_key=web_unblocker_api_key,  # Pass Web Unblocker for anti-bot
            web_unblocker_zone=web_unblocker_zone  # Pass Web Unblocker zone
        )
        
        # For preview, we need to wait longer for JavaScript to render
        # Product Hunt and similar React/Next.js sites need 10-15 seconds for hydration
        result = await fetcher.fetch(
            request.url,
            wait_for_selector=None,  # Don't wait for specific selector
            scroll_to_bottom=False,  # Don't scroll for preview
            click_load_more=None  # Don't click load more for preview
        )
        html = result.get('html', '')
        captured_json = result.get('captured_json', [])  # JSON from API responses
        fetch_method = result.get('fetch_method', 'unknown')
        fallback_reason = result.get('fallback_reason')
        error = result.get('error')
        
        # Check if browser rendering completely failed (no HTML returned)
        if fetch_method == 'browser_failed' or (not html and fetch_method == 'browser'):
            error_msg = fallback_reason or error or "Browser rendering failed"
            logger.error(f"[{tenant_id}] Browser rendering failed: {error_msg}")
            
            # For preview endpoint, try static HTML as fallback even if browser mode was forced
            # This ensures users can still see the page, even if JS didn't render
            logger.info(f"[{tenant_id}] Attempting static HTML fallback for preview...")
            try:
                from universal_scraper.core.html_fetcher import HTMLFetcher
                static_fetcher = HTMLFetcher(proxy_config=proxy_config)
                static_result = static_fetcher.fetch(request.url)
                static_html = static_result.get('html', '')
                
                if static_html and len(static_html) > 100:
                    logger.info(f"[{tenant_id}] Static HTML fallback successful ({len(static_html)} bytes)")
                    html = static_html
                    fetch_method = 'static_fallback'
                    fallback_reason = f"Browser rendering failed: {error_msg}. Showing static HTML version."
                else:
                    raise HTTPException(
                        status_code=400, 
                        detail=f"Failed to fetch page content. Browser rendering failed: {error_msg}. Static HTML fallback also failed."
                    )
            except Exception as static_err:
                logger.error(f"[{tenant_id}] Static HTML fallback also failed: {static_err}")
                raise HTTPException(
                    status_code=400, 
                    detail=f"Failed to fetch page content. Browser rendering failed: {error_msg}. Static HTML fallback also failed: {str(static_err)}"
                )
        
        # Check if browser rendering failed and we fell back to static HTML
        if fetch_method == 'static' or fetch_method == 'static_fallback':
            logger.warning(f"[{tenant_id}] Browser rendering failed, using static HTML fallback")
            if fallback_reason:
                logger.warning(f"[{tenant_id}] Fallback reason: {fallback_reason}")
        
        # Ensure HTML is fully rendered - add additional wait for JavaScript-heavy sites
        # Web Unblocker should return fully rendered HTML, but browser mode needs extra time
        if fetch_method == 'browser' or fetch_method == 'web_unblocker':
            # For preview, ensure we wait a bit longer for React/Next.js hydration
            # This is especially important for Product Hunt and similar sites
            import asyncio
            await asyncio.sleep(3)  # Additional 3s wait for JavaScript to fully render
        
        if not html or len(html) < 100:
            raise HTTPException(
                status_code=400, 
                detail=f"Failed to fetch page content. HTML is empty or too small ({len(html) if html else 0} bytes). Fetch method: {fetch_method}"
            )
        
        # Validate and clean HTML to prevent encoding issues
        # Check if HTML contains binary/corrupted data
        try:
            # Try to decode as UTF-8 if it's bytes
            if isinstance(html, bytes):
                html = html.decode('utf-8', errors='replace')
            
            # Check for binary/corrupted content indicators
            # If HTML contains too many non-printable characters, it's likely corrupted
            # But be more lenient - some sites have legitimate non-ASCII characters
            non_printable_count = sum(1 for c in html[:2000] if ord(c) < 32 and c not in '\n\r\t' and ord(c) != 9)
            if non_printable_count > 50:  # More lenient threshold
                logger.warning(f"[{tenant_id}] HTML appears corrupted (high non-printable count: {non_printable_count})")
                # Try to clean it more aggressively
                import re
                # Remove null bytes and other control characters except newlines/tabs/carriage returns
                html = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', html)
                # If still corrupted after cleaning, log but don't fail - return what we have
                if len(html) < 100:
                    logger.error(f"[{tenant_id}] HTML too corrupted after cleaning ({len(html)} bytes)")
                    raise HTTPException(status_code=400, detail="Page content appears corrupted or binary. Browser rendering may have failed.")
                else:
                    logger.info(f"[{tenant_id}] Cleaned HTML: {len(html)} bytes remaining")
            
            # Ensure HTML is valid UTF-8 string
            html = html.encode('utf-8', errors='replace').decode('utf-8')
            
        except (UnicodeDecodeError, UnicodeEncodeError) as e:
            logger.error(f"[{tenant_id}] HTML encoding error: {e}")
            raise HTTPException(status_code=400, detail=f"Page content encoding error: {str(e)}")
        
        soup = BeautifulSoup(html, 'html.parser')
        
        # JSON-FIRST: Use JSONDetector to find all JSON sources (like the scraper does)
        json_detector = JSONDetector()
        json_detection_result = json_detector.detect_and_extract(html, request.url, captured_json=captured_json)
        
        json_sources = []
        json_recommended = False
        json_confidence = 0.0
        
        if json_detection_result.get('json_found'):
            # Analyze JSON sources to determine if they're usable
            json_data = json_detection_result.get('data', [])
            json_source_types = json_detection_result.get('sources', [])
            
            # Check if JSON contains extractable data (arrays of objects)
            for idx, json_obj in enumerate(json_data):
                source_type = json_source_types[idx] if idx < len(json_source_types) else 'unknown'
                json_blob = json_obj.get('_data', {})
                
                # Check if it's an array of objects (likely extractable items)
                if isinstance(json_blob, list) and len(json_blob) > 0:
                    if isinstance(json_blob[0], dict):
                        # This looks like extractable data!
                        json_recommended = True
                        json_confidence = min(0.9, 0.5 + (len(json_blob) / 100))  # More items = higher confidence
                        
                        # Sample first item to show structure
                        sample_item = json_blob[0] if len(json_blob) > 0 else {}
                        json_sources.append({
                            'type': source_type,
                            'count': len(json_blob),
                            'sample_fields': list(sample_item.keys())[:10],  # First 10 fields
                            'preview': json_module.dumps(sample_item, indent=2)[:500] + '...' if len(str(sample_item)) > 500 else json_module.dumps(sample_item, indent=2),
                            'usable': True
                        })
                        break  # Found usable JSON, prioritize this
                
                # Also check for nested arrays in objects
                elif isinstance(json_blob, dict):
                    # Look for common array keys that might contain items
                    for key in ['items', 'products', 'results', 'data', 'list', 'articles', 'posts']:
                        if key in json_blob and isinstance(json_blob[key], list) and len(json_blob[key]) > 0:
                            json_recommended = True
                            json_confidence = 0.7
                            json_sources.append({
                                'type': source_type,
                                'count': len(json_blob[key]),
                                'sample_fields': list(json_blob[key][0].keys())[:10] if isinstance(json_blob[key][0], dict) else [],
                                'preview': json_module.dumps(json_blob[key][0], indent=2)[:500] + '...' if len(str(json_blob[key][0])) > 500 else json_module.dumps(json_blob[key][0], indent=2),
                                'usable': True,
                                'path': key
                            })
                            break
                    
                    if json_recommended:
                        break
            
            # If no usable JSON found, still show what was detected
            if not json_recommended and json_data:
                for idx, json_obj in enumerate(json_data):
                    source_type = json_source_types[idx] if idx < len(json_source_types) else 'unknown'
                    json_sources.append({
                        'type': source_type,
                        'preview': str(json_obj)[:500] + '...',
                        'usable': False
                    })
        
        # Detect repeating elements (potential items to extract)
        detected_elements = []
        
        # Common item container selectors
        item_selectors = [
            ('article', 'Article'),
            ('[role="article"]', 'Article (ARIA)'),
            ('[class*="item"]', 'Item'),
            ('[class*="card"]', 'Card'),
            ('[class*="product"]', 'Product'),
            ('[class*="listing"]', 'Listing'),
            ('[class*="result"]', 'Result'),
            ('[class*="post"]', 'Post'),
            ('li[class]', 'List Item'),
        ]
        
        for selector, field_name in item_selectors:
            try:
                elements = soup.select(selector)
                if len(elements) >= 3:  # At least 3 items = likely a list
                    # Get sample text from first element
                    sample = elements[0].get_text(strip=True)[:100] if elements else None
                    detected_elements.append(DetectedElement(
                        field=f"Container: {field_name}",
                        selector=selector,
                        sample_value=sample,
                        count=len(elements)
                    ))
            except:
                continue
        
        # Detect common field selectors within items
        field_selectors = {
            'title': ['h1', 'h2', 'h3', '[class*="title"]', '[class*="name"]'],
            'price': ['[class*="price"]', '[class*="cost"]', '[data-price]'],
            'description': ['[class*="description"]', '[class*="desc"]', 'p'],
            'image': ['img[src]', '[class*="image"] img', '[class*="photo"] img'],
            'rating': ['[class*="rating"]', '[class*="stars"]', '[class*="score"]'],
            'date': ['time', '[class*="date"]', '[datetime]'],
            'author': ['[class*="author"]', '[class*="user"]', '[class*="seller"]'],
            'location': ['[class*="location"]', '[class*="address"]'],
            'url': ['a[href]'],
        }
        
        for field, selectors in field_selectors.items():
            for selector in selectors:
                try:
                    elements = soup.select(selector)
                    if elements:
                        sample = elements[0].get_text(strip=True)[:100] if hasattr(elements[0], 'get_text') else str(elements[0].get('src', ''))[:100]
                        detected_elements.append(DetectedElement(
                            field=field,
                            selector=selector,
                            sample_value=sample,
                            count=len(elements)
                        ))
                        break  # Only add first matching selector per field
                except:
                    continue
        
        # Inject highlighting script into HTML
        highlight_script = '''
<script>
(function() {
    // ParaDocs Visual Preview Script
    window.paradocsPreview = {
        selectedElements: [],
        highlightColor: 'rgba(99, 102, 241, 0.3)',
        selectedColor: 'rgba(34, 197, 94, 0.5)',
        
        init: function() {
            // Add styles
            const style = document.createElement('style');
            style.textContent = `
                .paradocs-highlight {
                    outline: 2px dashed #6366f1 !important;
                    background-color: rgba(99, 102, 241, 0.1) !important;
                    cursor: pointer !important;
                    transition: all 0.2s ease !important;
                }
                .paradocs-highlight:hover {
                    outline: 3px solid #6366f1 !important;
                    background-color: rgba(99, 102, 241, 0.2) !important;
                }
                .paradocs-selected {
                    outline: 3px solid #22c55e !important;
                    background-color: rgba(34, 197, 94, 0.2) !important;
                }
                .paradocs-tooltip {
                    position: fixed;
                    background: #1f2937;
                    color: white;
                    padding: 8px 12px;
                    border-radius: 6px;
                    font-size: 12px;
                    z-index: 999999;
                    pointer-events: none;
                    max-width: 300px;
                    word-wrap: break-word;
                }
                .paradocs-container-highlight {
                    outline: 3px dashed #f59e0b !important;
                    background-color: rgba(245, 158, 11, 0.1) !important;
                }
            `;
            document.head.appendChild(style);
            
            // Create tooltip
            this.tooltip = document.createElement('div');
            this.tooltip.className = 'paradocs-tooltip';
            this.tooltip.style.display = 'none';
            document.body.appendChild(this.tooltip);
            
            // Highlight detectable elements
            this.highlightElements();
            
            // Listen for messages from parent
            window.addEventListener('message', (e) => this.handleMessage(e));
        },
        
        highlightElements: function() {
            // Make ALL elements clickable (not just specific selectors)
            // This allows users to click any element on the page
            document.addEventListener('click', (e) => {
                // Don't interfere with links and buttons
                if (e.target.tagName === 'A' || e.target.tagName === 'BUTTON') {
                    return;
                }
                
                // Prevent default navigation
                e.preventDefault();
                e.stopPropagation();
                
                // Handle click on any element
                this.handleClick(e, e.target);
            }, true); // Use capture phase to catch clicks early
            
            // Highlight common extractable elements for visual guidance
            const selectors = [
                'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                '[class*="title"]', '[class*="name"]',
                '[class*="price"]', '[class*="cost"]',
                '[class*="description"]', '[class*="desc"]',
                'img', 'picture',
                '[class*="rating"]', '[class*="score"]',
                'time', '[class*="date"]',
                '[class*="author"]', '[class*="user"]', '[class*="seller"]',
                '[class*="location"]', '[class*="address"]',
                'article', '[role="article"]',
                '[class*="item"]', '[class*="card"]', '[class*="product"]'
            ];
            
            selectors.forEach(sel => {
                try {
                    document.querySelectorAll(sel).forEach(el => {
                        if (!el.classList.contains('paradocs-highlight')) {
                            el.classList.add('paradocs-highlight');
                            el.addEventListener('mouseenter', (e) => this.showTooltip(e, el));
                            el.addEventListener('mouseleave', () => this.hideTooltip());
                        }
                    });
                } catch(e) {}
            });
        },
        
        handleClick: function(e, el) {
            e.preventDefault();
            e.stopPropagation();
            
            const selector = this.getSelector(el);
            const text = el.innerText?.substring(0, 100) || el.src || '';
            
            if (el.classList.contains('paradocs-selected')) {
                el.classList.remove('paradocs-selected');
                this.selectedElements = this.selectedElements.filter(s => s.selector !== selector);
            } else {
                el.classList.add('paradocs-selected');
                this.selectedElements.push({ selector, text, tagName: el.tagName });
            }
            
            // Send to parent
            window.parent.postMessage({
                type: 'paradocs-element-selected',
                element: { selector, text, tagName: el.tagName },
                allSelected: this.selectedElements
            }, '*');
        },
        
        showTooltip: function(e, el) {
            const selector = this.getSelector(el);
            const text = (el.innerText?.substring(0, 50) || el.src || '').trim();
            this.tooltip.innerHTML = `<strong>${el.tagName}</strong><br>${selector}<br><em>${text}...</em>`;
            this.tooltip.style.display = 'block';
            this.tooltip.style.left = e.clientX + 10 + 'px';
            this.tooltip.style.top = e.clientY + 10 + 'px';
        },
        
        hideTooltip: function() {
            this.tooltip.style.display = 'none';
        },
        
        getSelector: function(el) {
            // Generate a robust CSS selector
            // Priority: ID > unique class combination > data attributes > nth-child > tag
            
            // 1. ID selector (most specific)
            if (el.id) {
                return '#' + el.id;
            }
            
            // 2. Try to find unique class combination
            if (el.className && typeof el.className === 'string') {
                const classes = el.className.split(' ')
                    .filter(c => c && !c.startsWith('paradocs') && c.length > 0);
                
                if (classes.length > 0) {
                    // Try with tag + first class
                    const selector1 = el.tagName.toLowerCase() + '.' + classes[0];
                    const matches1 = document.querySelectorAll(selector1).length;
                    
                    if (matches1 === 1) {
                        return selector1;
                    }
                    
                    // Try with tag + multiple classes
                    if (classes.length > 1) {
                        const selector2 = el.tagName.toLowerCase() + '.' + classes.slice(0, 2).join('.');
                        const matches2 = document.querySelectorAll(selector2).length;
                        if (matches2 <= 5) { // Reasonable number of matches
                            return selector2;
                        }
                    }
                }
            }
            
            // 3. Try data attributes
            if (el.hasAttribute('data-testid')) {
                return '[data-testid="' + el.getAttribute('data-testid') + '"]';
            }
            if (el.hasAttribute('data-id')) {
                return '[data-id="' + el.getAttribute('data-id') + '"]';
            }
            
            // 4. Use nth-child as fallback
            const parent = el.parentElement;
            if (parent) {
                const siblings = Array.from(parent.children).filter(c => c.tagName === el.tagName);
                const index = siblings.indexOf(el);
                if (index >= 0 && siblings.length > 1) {
                    return el.tagName.toLowerCase() + ':nth-child(' + (index + 1) + ')';
                }
            }
            
            // 5. Last resort: tag name
            return el.tagName.toLowerCase();
        },
        
        handleMessage: function(e) {
            if (e.data.type === 'paradocs-highlight-selector') {
                document.querySelectorAll(e.data.selector).forEach(el => {
                    el.classList.add('paradocs-container-highlight');
                });
            }
        }
    };
    
    // Initialize when DOM is ready
    if (document.readyState === 'loading') {
        document.addEventListener('DOMContentLoaded', () => window.paradocsPreview.init());
    } else {
        window.paradocsPreview.init();
    }
})();
</script>
'''
        
        # Inject script before </body>
        if '</body>' in html:
            html_with_script = html.replace('</body>', highlight_script + '</body>')
        else:
            html_with_script = html + highlight_script
        
        # Get page info
        title = soup.find('title')
        page_info = {
            'title': title.get_text() if title else 'Unknown',
            'url': request.url,
            'html_size': len(html),
            'element_count': len(soup.find_all()),
            'has_json': len(json_sources) > 0
        }
        
        # Determine recommended extraction mode
        extraction_mode = "json" if json_recommended else "browser"
        
        # Check if browser rendering failed
        browser_rendering_failed = fetch_method in ('static', 'static_fallback')
        
        return PreviewResponse(
            success=True,
            html=html_with_script,
            detected_elements=detected_elements,
            json_sources=json_sources if json_sources else None,
            json_recommended=json_recommended,
            json_confidence=json_confidence,
            extraction_mode=extraction_mode,
            fetch_method=fetch_method,
            browser_rendering_failed=browser_rendering_failed,
            fallback_reason=fallback_reason,
            page_info=page_info,
            unblocker_log=result.get('unblocker_log', [])
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[{tenant_id}] Preview failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Preview failed: {str(e)}")


@app.post("/scrape")
async def scrape_endpoint(
    request: ScrapeRequest,
    tenant_id: str = Depends(get_tenant_id),
    tenant_config: Dict = Depends(get_tenant_context),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Scrape a single URL (Multi-tenant aware)
    
    Requires authentication (Bearer token or X-API-Key header)
    """
    import time as time_module
    start_time = time_module.time()
    
    # Get API key from header or environment
    api_key = x_api_key or os.getenv("OPENAI_API_KEY") or os.getenv("GEMINI_API_KEY")
    
    if not api_key:
        raise HTTPException(
            status_code=401,
            detail="API key required. Provide X-API-Key header or set OPENAI_API_KEY/GEMINI_API_KEY environment variable"
        )
    
    try:
        logger.info(f"[{tenant_id}] Scraping URL: {request.url}")
        logger.info(f"[{tenant_id}] Fields: {request.fields if request.fields else 'AUTO-EXTRACT'}")
        
        # Check rate limits
        rate_limiter = get_rate_limiter()
        await rate_limiter.check_rate_limit(tenant_id, tenant_config, request.url)
        
        # Initialize tenant-aware cache
        tenant_cache = TenantCache(tenant_id, get_redis_cache())
        
        # Check tenant-specific execution cache first
        cached_result = await tenant_cache.get_execution_cache(request.url, request.fields)
        if cached_result:
            logger.info(f"[{tenant_id}] Cache HIT (execution): {request.url}")
            execution_time_ms = int((time_module.time() - start_time) * 1000)
            
            # Track usage (cache hit)
            usage_tracker = get_usage_tracker()
            await usage_tracker.track_request(
                tenant_id=tenant_id,
                endpoint="/scrape",
                url=request.url,
                items_extracted=len(cached_result.get("data", [])),
                cache_hit=True,
                execution_time_ms=execution_time_ms
            )
            
            return {
                "success": True,
                "data": cached_result.get("data", []),
                "metadata": {
                    **cached_result.get("metadata", {}),
                    "cache_hit": True,
                    "tenant_id": tenant_id,
                },
                "source": cached_result.get("metadata", {}).get("source", "cache")
            }
        
        logger.info(f"[{tenant_id}] Cache MISS (execution): {request.url}")
        logger.info(f"[{tenant_id}] Mode: {request.mode}")
        logger.info(f"[{tenant_id}] Scroll to bottom: {request.scroll_to_bottom}")
        logger.info(f"[{tenant_id}] Wait for selector: {request.wait_for_selector}")
        if request.proxy_config:
            logger.info(f"[{tenant_id}] Using proxy: {request.proxy_config.get('server', 'unknown')}")
        
        # Get Redis cache for tenant-aware caching
        redis_cache = get_redis_cache()
        scraper = get_scraper(
            api_key,
            request.mode,
            request.proxy_config,
            redis_cache,
            browser_timeout=request.browser_timeout or 120000  # Default 120s, configurable
        )
        
        result = await scraper.scrape(
            url=request.url,
            fields=request.fields,
            target=request.target,
            html=request.html,
            force_html=request.force_html,
            force_generate=request.force_generate,
            scroll_to_bottom=request.scroll_to_bottom,
            click_load_more=request.click_load_more,
            wait_for_selector=request.wait_for_selector
        )
        
        execution_time_ms = int((time_module.time() - start_time) * 1000)
        
        # Cache result (tenant-specific)
        await tenant_cache.set_execution_cache(
            request.url,
            request.fields,
            result,
            ttl=tenant_config.get("cache_ttl", 3600)
        )
        
        # Track usage (cache miss)
        usage_tracker = get_usage_tracker()
        await usage_tracker.track_request(
            tenant_id=tenant_id,
            endpoint="/scrape",
            url=request.url,
            items_extracted=len(result.get("data", [])),
            cache_hit=False,
            execution_time_ms=execution_time_ms
        )
        
        # Check if cache was stored (code cache or Direct LLM cache)
        cache_stored = False
        cache_type = None
        if result.get("metadata", {}).get("code_cached") is False:
            # Code was generated and should be cached
            cache_stored = True
            cache_type = "code"
        elif result.get("source") == "direct_llm":
            # Direct LLM was used and should be cached
            cache_stored = True
            cache_type = "direct_llm"
        
        return {
            "success": True,
            "data": result.get("data", []),
            "metadata": {
                **result.get("metadata", {}),
                "cache_hit": False,
                "cache_stored": cache_stored,
                "cache_type": cache_type,
                "tenant_id": tenant_id,
            },
            "source": result.get("source", "unknown")
        }
    
    except HTTPException:
        # Re-raise HTTP exceptions (rate limit, auth, etc.)
        raise
    except Exception as e:
        error_str = str(e).lower()
        error_message = str(e)
        
        # Detect blocking/timeout issues
        is_timeout = "timeout" in error_str or "timed out" in error_str
        is_blocked = any([
            "blocked" in error_str,
            "access denied" in error_str,
            "forbidden" in error_str,
            "403" in error_str,
            "cloudflare" in error_str,
            "captcha" in error_str,
            "bot detection" in error_str,
            "you've been blocked" in error_str,
            "network security" in error_str
        ])
        
        # Check if proxy was used
        proxy_used = request.proxy_config is not None and bool(request.proxy_config)
        is_web_unlocker = proxy_used and (
            "web_unlocker" in str(request.proxy_config.get("username", "")).lower() or
            "unlocker" in str(request.proxy_config.get("server", "")).lower()
        )
        
        # Determine recommendation
        recommendation = None
        if is_timeout or is_blocked:
            if not proxy_used:
                recommendation = {
                    "type": "proxy_required",
                    "message": "This site appears to be blocking automated requests. Try using a residential proxy or Web Unlocker.",
                    "action": "configure_proxy",
                    "severity": "high"
                }
            elif not is_web_unlocker:
                recommendation = {
                    "type": "web_unlocker_recommended",
                    "message": "This site uses advanced anti-bot protection. Web Unlocker is recommended for better success rates.",
                    "action": "upgrade_to_web_unlocker",
                    "severity": "medium"
                }
            else:
                recommendation = {
                    "type": "advanced_blocking",
                    "message": "Site is using advanced anti-bot measures. You may need to adjust settings or try again later.",
                    "action": "retry_or_contact_support",
                    "severity": "high"
                }
        
        logger.error(f"[{tenant_id}] Scraping failed: {error_message}", exc_info=True)
        
        # Return error with helpful metadata
        # FastAPI will serialize dict details as JSON
        error_detail = {
            "error": error_message,
            "is_timeout": is_timeout,
            "is_blocked": is_blocked,
            "proxy_used": proxy_used,
            "is_web_unlocker": is_web_unlocker,
            "recommendation": recommendation
        }
        
        raise HTTPException(
            status_code=500,
            detail=error_detail
        )

@app.post("/crawl")
async def crawl_endpoint(
    request: CrawlRequest,
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Crawl multiple URLs
    
    Requires X-API-Key header with OpenAI/Gemini/Claude API key
    """
    # Get API key from header or environment
    api_key = x_api_key or os.getenv("OPENAI_API_KEY") or os.getenv("GEMINI_API_KEY")
    
    if not api_key:
        raise HTTPException(
            status_code=401,
            detail="API key required. Provide X-API-Key header or set OPENAI_API_KEY/GEMINI_API_KEY environment variable"
        )
    
    try:
        logger.info(f"Crawling {len(request.start_urls)} URLs")
        
        scraper = get_scraper(api_key, "hybrid")
        
        all_results = []
        for url in request.start_urls:
            try:
                result = await scraper.scrape(
                    url=url,
                    fields=request.fields
                )
                all_results.append({
                    "url": url,
                    "success": True,
                    "data": result.get("data", []),
                    "metadata": result.get("metadata", {})
                })
            except Exception as e:
                logger.error(f"Failed to scrape {url}: {str(e)}")
                all_results.append({
                    "url": url,
                    "success": False,
                    "error": str(e)
                })
        
        return {
            "success": True,
            "results": all_results,
            "total_urls": len(request.start_urls),
            "successful": sum(1 for r in all_results if r.get("success"))
        }
    
    except Exception as e:
        logger.error(f"Crawling failed: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Crawling failed: {str(e)}"
        )

@app.post("/document-processing/extract", response_model=DocumentProcessingResponse)
async def document_processing_endpoint(
    file: UploadFile = File(...),
    fields: str = Form("[]"),  # JSON string of fields array
    use_ocr: bool = Form(False),
    max_pages: Optional[int] = Form(None),
    context: Optional[str] = Form(None),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Process document and extract structured data
    
    Requires X-API-Key header with OpenAI/Gemini/Claude API key
    """
    # Get API key from header or environment
    api_key = x_api_key or os.getenv("OPENAI_API_KEY") or os.getenv("GEMINI_API_KEY")
    
    if not api_key:
        raise HTTPException(
            status_code=401,
            detail="API key required. Provide X-API-Key header or set OPENAI_API_KEY/GEMINI_API_KEY environment variable"
        )
    
    try:
        import json
        
        # Parse fields
        try:
            fields_list = json.loads(fields) if fields else []
        except json.JSONDecodeError:
            fields_list = []
        
        logger.info(f"Processing document: {file.filename}")
        logger.info(f"Fields: {fields_list if fields_list else 'AUTO-EXTRACT'}")
        logger.info(f"OCR: {use_ocr}")
        
        # Extract text from document
        text = await extract_text_from_file(file, use_ocr)
        
        # Limit pages if specified (for PDF)
        if max_pages and file.filename.lower().endswith('.pdf'):
            lines = text.split('\n')
            # Rough estimate: ~50 lines per page
            lines_per_page = 50
            max_lines = max_pages * lines_per_page
            text = '\n'.join(lines[:max_lines])
        
        logger.info(f"Extracted {len(text):,} characters from document")
        
        # Use DirectLLMExtractor to extract structured data (with Redis cache)
        redis_cache = get_redis_cache()
        extractor = get_extractor(api_key, redis_cache=redis_cache)
        
        # Convert text to HTML-like format for extractor (it expects HTML)
        # Wrap in basic HTML structure
        html_content = f"<html><body><pre>{text}</pre></body></html>"
        
        # Generate cache key for document (based on filename + fields)
        # Documents are cached by filename pattern (e.g., "invoice.pdf" pattern)
        import hashlib
        filename_hash = hashlib.md5(file.filename.encode()).hexdigest()[:8]
        fields_str = ','.join(sorted(fields_list)) if fields_list else 'auto'
        fields_hash = hashlib.md5(fields_str.encode()).hexdigest()[:8]
        document_cache_key = f"document_{filename_hash}_{fields_hash}"
        
        # Extract data using LLM (will cache automatically via DirectLLMExtractor)
        # Pass a URL-like identifier for cache key generation
        document_url = f"document://{file.filename}"
        extracted_data = await extractor.extract(
            html=html_content,
            fields=fields_list if fields_list else ['text', 'title', 'metadata'],
            context=context,
            url=document_url  # Pass URL for cache key generation
        )
        
        logger.info(f"Extracted {len(extracted_data)} items from document (cached for future similar documents)")
        
        return {
            "success": True,
            "data": extracted_data,
            "metadata": {
                "filename": file.filename,
                "file_size": len(text),
                "items_extracted": len(extracted_data),
                "use_ocr": use_ocr,
                "cache_key": document_cache_key,  # Include cache key in response
            }
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document processing failed: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Document processing failed: {str(e)}"
        )

@app.post("/api/v1/proxy/test", response_model=ProxyTestResponse)
async def test_proxy(request: ProxyTestRequest):
    """
    Test proxy connection
    
    Supports:
    - Bright Data (brightdata)
    - Oxylabs (oxylabs)
    - ScraperAPI (scraperapi)
    - Custom proxies
    """
    try:
        import requests
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        
        # Build proxy config based on provider
        proxy_config = {}
        
        if request.provider == 'brightdata':
            # Bright Data format
            server_raw = request.server or (request.externalProxy.get('server') if request.externalProxy else 'brd.superproxy.io:33335')
            username = request.username or (request.externalProxy.get('username') if request.externalProxy else '')
            password = request.password or (request.externalProxy.get('password') if request.externalProxy else '')
            
            # Parse server - handle both "host:port" and "host,port" formats
            # Fix common mistake: comma instead of colon
            if ',' in server_raw and ':' not in server_raw:
                server_raw = server_raw.replace(',', ':')
            
            # Ensure server has port
            if ':' not in server_raw:
                server_raw = f"{server_raw}:33335"  # Default Bright Data port
            
            # Extract host and port
            if '://' in server_raw:
                # Remove protocol if present
                server_raw = server_raw.split('://', 1)[1]
            
            server = server_raw  # Now in format "host:port"
            
            # Add country to username if provided (Bright Data format: username-country-US)
            if request.country or (request.externalProxy and request.externalProxy.get('country')):
                country = request.country or request.externalProxy.get('country')
                if country and f'-country-{country}' not in username:
                    username = f"{username}-country-{country}"
            
            proxy_config = {
                'server': f'http://{server}',
                'username': username,
                'password': password
            }
            
            # Build proxy URL - server is already in "host:port" format
            proxy_url = f"http://{username}:{password}@{server}"
            
            # Test with Bright Data's test endpoint
            test_url = "https://geo.brdtest.com/welcome.txt?product=resi&method=native"
            
            logger.info(f"Testing Bright Data proxy: {server}")
            
            response = requests.get(
                test_url,
                proxies={
                    'http': proxy_url,
                    'https': proxy_url
                },
                timeout=30,
                verify=False  # Bright Data test endpoint uses self-signed cert
            )
            
            if response.status_code == 200:
                # Extract IP from response
                ip_info = response.text
                return ProxyTestResponse(
                    success=True,
                    message=f"Proxy connection successful! Response: {ip_info[:200]}",
                    ip=ip_info.split('IP:')[1].split()[0] if 'IP:' in ip_info else None
                )
            else:
                return ProxyTestResponse(
                    success=False,
                    message=f"Proxy test failed with status {response.status_code}: {response.text[:200]}"
                )
        
        elif request.provider == 'oxylabs':
            # Oxylabs format
            server = request.server or request.externalProxy.get('server') if request.externalProxy else 'pr.oxylabs.io:7777'
            username = request.username or request.externalProxy.get('username') if request.externalProxy else ''
            password = request.password or request.externalProxy.get('password') if request.externalProxy else ''
            
            # Add country to username if provided
            if request.country or (request.externalProxy and request.externalProxy.get('country')):
                country = request.country or request.externalProxy.get('country')
                if country and f'-country-{country}' not in username:
                    username = f"{username}-country-{country}"
            
            proxy_url = f"http://{username}:{password}@{server.replace('http://', '').replace('https://', '')}"
            
            # Test with a simple HTTP request
            test_url = "http://httpbin.org/ip"
            
            logger.info(f"Testing Oxylabs proxy: {server}")
            
            response = requests.get(
                test_url,
                proxies={
                    'http': proxy_url,
                    'https': proxy_url
                },
                timeout=30
            )
            
            if response.status_code == 200:
                ip_data = response.json()
                return ProxyTestResponse(
                    success=True,
                    message="Proxy connection successful!",
                    ip=ip_data.get('origin', '').split(',')[0]
                )
            else:
                return ProxyTestResponse(
                    success=False,
                    message=f"Proxy test failed with status {response.status_code}"
                )
        
        elif request.provider == 'web_unlocker' or request.provider == 'webunblocker':
            # Web Unblocker test - uses Bright Data Web Unblocker API
            web_unblocker = request.webUnblocker or {}
            api_key = web_unblocker.get('apiKey') or web_unblocker.get('api_key')
            zone = web_unblocker.get('zone', 'web_unlocker1')
            
            if not api_key:
                return ProxyTestResponse(
                    success=False,
                    message="Web Unblocker API key is required"
                )
            
            logger.info(f"Testing Web Unblocker: zone={zone}")
            
            try:
                # Use Web Unblocker to fetch a test page (like Product Hunt which has Cloudflare)
                # This will verify the API key works and can bypass anti-bot protection
                from universal_scraper.core.web_unblocker_fetcher import WebUnblockerFetcher
                
                fetcher = WebUnblockerFetcher(
                    api_key=api_key,
                    zone=zone
                )
                
                # Test with a Cloudflare-protected site (Product Hunt)
                test_url = "https://www.producthunt.com/"
                
                # Use await directly (we're already in an async function)
                result = await fetcher.fetch_async(test_url)
                
                html = result.get('html', '')
                
                if html and len(html) > 1000:
                    # Check if we got actual content (not Cloudflare challenge)
                    if 'verify you are human' in html.lower() or 'just a moment' in html.lower():
                        return ProxyTestResponse(
                            success=False,
                            message="Web Unblocker connected but Cloudflare challenge detected. Check API key and zone configuration."
                        )
                    
                    # Success - got real content
                    return ProxyTestResponse(
                        success=True,
                        message=f"Web Unblocker connection successful! Successfully bypassed Cloudflare protection. Retrieved {len(html)} bytes of content.",
                        ip=None  # Web Unblocker doesn't expose IP
                    )
                else:
                    return ProxyTestResponse(
                        success=False,
                        message=f"Web Unblocker test returned insufficient content ({len(html)} bytes). Check API key and zone."
                    )
                    
            except Exception as e:
                logger.error(f"Web Unblocker test failed: {e}", exc_info=True)
                return ProxyTestResponse(
                    success=False,
                    message=f"Web Unblocker test failed: {str(e)}"
                )
        
        elif request.provider == 'custom' or request.provider == 'scraperapi':
            # Custom/ScraperAPI format
            server = request.server or request.externalProxy.get('server') if request.externalProxy else ''
            username = request.username or request.externalProxy.get('username') if request.externalProxy else ''
            password = request.password or request.externalProxy.get('password') if request.externalProxy else ''
            
            if not server or not username or not password:
                return ProxyTestResponse(
                    success=False,
                    message="Missing required proxy fields: server, username, and password are required"
                )
            
            proxy_url = f"http://{username}:{password}@{server.replace('http://', '').replace('https://', '')}"
            
            # Test with a simple HTTP request
            test_url = "http://httpbin.org/ip"
            
            logger.info(f"Testing custom proxy: {server}")
            
            response = requests.get(
                test_url,
                proxies={
                    'http': proxy_url,
                    'https': proxy_url
                },
                timeout=30
            )
            
            if response.status_code == 200:
                ip_data = response.json()
                return ProxyTestResponse(
                    success=True,
                    message="Proxy connection successful!",
                    ip=ip_data.get('origin', '').split(',')[0]
                )
            else:
                return ProxyTestResponse(
                    success=False,
                    message=f"Proxy test failed with status {response.status_code}"
                )
        
        else:
            return ProxyTestResponse(
                success=False,
                message=f"Unsupported proxy provider: {request.provider}"
            )
    
    except requests.exceptions.ProxyError as e:
        logger.error(f"Proxy error: {e}")
        return ProxyTestResponse(
            success=False,
            message=f"Proxy connection failed: {str(e)}"
        )
    except requests.exceptions.Timeout:
        return ProxyTestResponse(
            success=False,
            message="Proxy test timed out. Check your proxy configuration and network connection."
        )
    except Exception as e:
        logger.error(f"Proxy test error: {e}", exc_info=True)
        return ProxyTestResponse(
            success=False,
            message=f"Proxy test failed: {str(e)}"
        )

@app.post("/api/v1/web-unblocker/test")
async def test_web_unblocker(
    request: dict,
    tenant_id: str = Depends(get_tenant_id)
):
    """
    Test Web Unblocker connection
    
    Request body:
    {
        "apiKey": "your-api-key",
        "zone": "web_unlocker1"
    }
    """
    try:
        api_key = request.get("apiKey")
        zone = request.get("zone", "web_unlocker1")
        
        if not api_key:
            return {"success": False, "message": "API key is required"}
        
        # Use WebUnblockerFetcher which handles both API key and proxy credentials
        from universal_scraper.core.web_unblocker_fetcher import WebUnblockerFetcher
        
        fetcher = WebUnblockerFetcher(
            api_key=api_key,
            zone=zone,
            timeout=30
        )
        
        # Test Web Unblocker by making a simple request to a reliable test URL
        # We use the fetch_async method which is already available
        test_url = "https://lumtest.com/myip.json"
        result = await fetcher.fetch_async(test_url)
        
        if result.get('success'):
            # Try to parse IP from result
            import json
            try:
                ip_data = json.loads(result.get('html', '{}'))
                ip = ip_data.get('ip', 'unknown')
                return {
                    "success": True,
                    "message": f"Web Unblocker connection successful! IP: {ip}"
                }
            except:
                return {
                    "success": True,
                    "message": "Web Unblocker connection successful!"
                }
        else:
            return {
                "success": False,
                "message": f"Web Unblocker test failed: {result.get('error', 'Unknown error')}"
            }
            
    except Exception as e:
        logger.error(f"[{tenant_id}] Web Unblocker test failed: {e}")
        return {
            "success": False,
            "message": f"Web Unblocker test failed: {str(e)}"
        }

@app.get("/api/v1/cache/check")
async def check_cache_status(
    url: str,
    fields: Optional[str] = None,
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Check if a URL/domain has cached extraction pattern
    
    Args:
        url: URL to check
        fields: Comma-separated list of fields (optional)
    """
    api_key = x_api_key or os.getenv("OPENAI_API_KEY") or os.getenv("GEMINI_API_KEY")
    
    if not api_key:
        raise HTTPException(
            status_code=401,
            detail="API key required"
        )
    
    try:
        import asyncio
        from urllib.parse import urlparse
        parsed_url = urlparse(url)
        domain = parsed_url.netloc
        
        # Parse fields if provided
        fields_list = []
        if fields:
            fields_list = [f.strip() for f in fields.split(',') if f.strip()]
        
        scraper = get_scraper(api_key, "hybrid", redis_cache=get_redis_cache())
        
        # Check if domain has cached patterns (code cache) - async with timeout
        try:
            patterns = await asyncio.wait_for(
                scraper.list_cached_patterns(domain),
                timeout=5.0  # 5 second timeout
            )
        except asyncio.TimeoutError:
            logger.warning(f"Cache check timeout for domain {domain}")
            patterns = []
        
        # Also check DirectLLM cache if available
        direct_llm_cached = False
        direct_llm_cache_age = None
        if hasattr(scraper, 'direct_llm_extractor') and scraper.direct_llm_extractor:
            if hasattr(scraper.direct_llm_extractor, 'result_cache') and scraper.direct_llm_extractor.result_cache:
                try:
                    # Generate cache key to check (matches DirectLLM cache key format)
                    fields_str = ','.join(sorted(fields_list)) if fields_list else ''
                    fields_hash = hashlib.md5(fields_str.encode()).hexdigest()[:8]
                    cache_key = f"direct_llm_{domain.replace('.', '_')}_{fields_hash}"
                    
                    cached_result = await scraper.direct_llm_extractor.result_cache.backend.get(cache_key)
                    if cached_result:
                        direct_llm_cached = True
                        timestamp = cached_result.get('timestamp', 0)
                        if timestamp:
                            direct_llm_cache_age = time.time() - timestamp
                except Exception as e:
                    logger.debug(f"DirectLLM cache check failed: {e}")
        
        # Check if specific fields pattern exists
        is_cached = False
        cache_age = None
        matched_pattern = None
        
        # Check code cache patterns
        if patterns:
            # If fields specified, check for exact match
            if fields_list:
                for pattern in patterns:
                    pattern_fields = pattern.get('fields', [])
                    if set(pattern_fields) == set(fields_list):
                        is_cached = True
                        matched_pattern = pattern
                        # Calculate cache age
                        created_at = pattern.get('created_at', 0)
                        if created_at:
                            cache_age = time.time() - created_at
                        break
            else:
                # If no fields specified, any pattern for domain counts as cached
                is_cached = True
                matched_pattern = patterns[0]
                created_at = patterns[0].get('created_at', 0)
                if created_at:
                    cache_age = time.time() - created_at
        
        # If code cache not found, check DirectLLM cache
        if not is_cached and direct_llm_cached:
            is_cached = True
            cache_age = direct_llm_cache_age
        
        return {
            "success": True,
            "is_cached": is_cached,
            "cache_age": cache_age,
            "domain": domain,
            "pattern": matched_pattern
        }
    except Exception as e:
        logger.error(f"Cache check failed: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Cache check failed: {str(e)}"
        )

@app.get("/api/v1/usage/stats")
async def get_usage_stats(
    tenant_id: str = Depends(get_tenant_id),
    date: Optional[str] = None
):
    """
    Get usage statistics for tenant
    """
    try:
        usage_tracker = get_usage_tracker()
        stats = await usage_tracker.get_usage_stats(tenant_id, date)
        
        rate_limiter = get_rate_limiter()
        rate_status = await rate_limiter.get_rate_limit_status(tenant_id)
        
        return {
            "success": True,
            "tenant_id": tenant_id,
            "usage": stats,
            "rate_limits": rate_status
        }
    except Exception as e:
        logger.error(f"Usage stats failed: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Usage stats failed: {str(e)}"
        )

@app.get("/api/v1/cache/patterns")
async def list_cached_patterns(
    domain: Optional[str] = None,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    List cached extraction patterns by domain
    """
    api_key = x_api_key or os.getenv("OPENAI_API_KEY") or os.getenv("GEMINI_API_KEY")
    
    if not api_key:
        raise HTTPException(
            status_code=401,
            detail="API key required"
        )
    
    try:
        scraper = get_scraper(api_key, "hybrid", redis_cache=get_redis_cache())
        patterns = await scraper.list_cached_patterns(domain)
        domains = await scraper.get_cached_domains()
        
        logger.info(f"Cache patterns query: found {len(patterns)} patterns, {len(domains)} domains")
        
        # Also include Direct LLM cache patterns if available
        direct_llm_patterns = []
        if hasattr(scraper, 'direct_llm_extractor') and scraper.direct_llm_extractor:
            if hasattr(scraper.direct_llm_extractor, 'result_cache') and scraper.direct_llm_extractor.result_cache:
                try:
                    backend = scraper.direct_llm_extractor.result_cache.backend
                    
                    # Handle different backend types
                    if hasattr(backend, 'list_keys'):
                        # UnifiedPatternCache backend (LocalFileCache, ApifyKVCache, or RedisCacheBackend)
                        cache_keys = await backend.list_keys(prefix="direct_llm_")
                    elif hasattr(backend, 'redis_cache') and backend.redis_cache:
                        # RedisCacheBackend - use its list_keys method
                        cache_keys = await backend.list_keys(prefix="direct_llm_")
                    elif hasattr(backend, 'redis_client') and backend.redis_client:
                        # Direct Redis client - use scan_iter
                        cache_keys = []
                        async for key in backend.redis_client.scan_iter(match="direct_llm_*"):
                            cache_keys.append(key)
                    else:
                        # Try Redis cache directly if available
                        redis_cache = get_redis_cache()
                        if redis_cache and redis_cache.redis_client:
                            cache_keys = []
                            async for key in redis_cache.redis_client.scan_iter(match="direct_llm_*"):
                                cache_keys.append(key)
                        else:
                            cache_keys = []
                    
                    logger.info(f"Found {len(cache_keys)} Direct LLM cache keys")
                    
                    for key in cache_keys:
                        try:
                            # Get cached data
                            if hasattr(backend, 'get'):
                                cached_data = await backend.get(key)
                            elif hasattr(backend, 'redis_cache') and backend.redis_cache:
                                # RedisCacheBackend - use redis_cache.get
                                cached_data = await backend.redis_cache.get(key)
                            elif hasattr(backend, 'redis_client') and backend.redis_client:
                                # Direct Redis client
                                data = await backend.redis_client.get(key)
                                cached_data = json.loads(data) if data else None
                            else:
                                # Fallback to direct Redis cache
                                redis_cache = get_redis_cache()
                                cached_data = await redis_cache.get(key) if redis_cache else None
                            
                            if cached_data:
                                # Extract domain from cache data (stored in cache_entry) - PREFERRED
                                cached_domain = cached_data.get('domain')
                                
                                # Fallback: extract from key if domain not in cache entry
                                # Key format: direct_llm_{domain_normalized}_{fields_hash}
                                # or: direct_llm_{structure_hash}_{fields_hash}
                                if not cached_domain:
                                    key_without_prefix = key.replace('direct_llm_', '')
                                    key_parts = key_without_prefix.split('_')
                                    if len(key_parts) >= 2:
                                        # Check if first part looks like a domain (has letters) vs hash (hex)
                                        first_part = key_parts[0]
                                        # If it contains letters (not just hex), it's likely a domain
                                        if any(c.isalpha() and c not in 'abcdef' for c in first_part.lower()):
                                            cached_domain = first_part.replace('_', '.')
                                
                                # Filter by domain if specified
                                if domain and cached_domain and cached_domain != domain:
                                    continue
                                
                                direct_llm_patterns.append({
                                    'cache_key': key,
                                    'domain': cached_domain or 'unknown',
                                    'fields': cached_data.get('fields', []),
                                    'url': cached_data.get('url', ''),
                                    'cache_type': 'direct_llm',
                                    'created_at': cached_data.get('timestamp', 0),
                                    'structure_hash': cached_data.get('structure_hash', ''),
                                    'item_count': cached_data.get('item_count', len(cached_data.get('items', [])))
                                })
                        except Exception as e:
                            logger.warning(f"Failed to process cache key {key}: {e}")
                            continue
                            
                except Exception as e:
                    logger.error(f"Failed to list Direct LLM cache patterns: {e}", exc_info=True)
        
        # Combine code cache and Direct LLM cache patterns
        all_patterns = patterns + direct_llm_patterns
        
        # Get unique domains from all patterns
        all_domains = set()
        for pattern in all_patterns:
            if pattern.get('domain'):
                all_domains.add(pattern['domain'])
        
        return {
            "success": True,
            "patterns": all_patterns,
            "domains": sorted(list(all_domains)),
            "total_patterns": len(all_patterns)
        }
    except Exception as e:
        logger.error(f"Failed to list cached patterns: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to list cached patterns: {str(e)}"
        )

@app.get("/api/v1/cache/export")
async def export_cache(
    domain: Optional[str] = None,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Export cached patterns as JSON file for download/sharing
    """
    api_key = x_api_key or os.getenv("OPENAI_API_KEY") or os.getenv("GEMINI_API_KEY")
    
    if not api_key:
        raise HTTPException(
            status_code=401,
            detail="API key required"
        )
    
    try:
        scraper = get_scraper(api_key, "hybrid", redis_cache=get_redis_cache())
        patterns = await scraper.list_cached_patterns(domain)
        domains = await scraper.get_cached_domains()
        
        # Also include Direct LLM cache patterns (same logic as list endpoint)
        direct_llm_patterns = []
        if hasattr(scraper, 'direct_llm_extractor') and scraper.direct_llm_extractor:
            if hasattr(scraper.direct_llm_extractor, 'result_cache') and scraper.direct_llm_extractor.result_cache:
                try:
                    backend = scraper.direct_llm_extractor.result_cache.backend
                    
                    if hasattr(backend, 'list_keys'):
                        cache_keys = await backend.list_keys(prefix="direct_llm_")
                    elif hasattr(backend, 'redis_cache') and backend.redis_cache:
                        cache_keys = await backend.list_keys(prefix="direct_llm_")
                    elif hasattr(backend, 'redis_client') and backend.redis_client:
                        cache_keys = []
                        async for key in backend.redis_client.scan_iter(match="direct_llm_*"):
                            cache_keys.append(key)
                    else:
                        redis_cache = get_redis_cache()
                        if redis_cache and redis_cache.redis_client:
                            cache_keys = []
                            async for key in redis_cache.redis_client.scan_iter(match="direct_llm_*"):
                                cache_keys.append(key)
                        else:
                            cache_keys = []
                    
                    for key in cache_keys:
                        try:
                            if hasattr(backend, 'get'):
                                cached_data = await backend.get(key)
                            elif hasattr(backend, 'redis_cache') and backend.redis_cache:
                                cached_data = await backend.redis_cache.get(key)
                            elif hasattr(backend, 'redis_client') and backend.redis_client:
                                data = await backend.redis_client.get(key)
                                cached_data = json.loads(data) if data else None
                            else:
                                redis_cache = get_redis_cache()
                                cached_data = await redis_cache.get(key) if redis_cache else None
                            
                            if cached_data:
                                cached_domain = cached_data.get('domain')
                                if domain and cached_domain != domain:
                                    continue
                                
                                direct_llm_patterns.append({
                                    'cache_key': key,
                                    'domain': cached_domain or 'unknown',
                                    'fields': cached_data.get('fields', []),
                                    'url': cached_data.get('url', ''),
                                    'cache_type': 'direct_llm',
                                    'created_at': cached_data.get('timestamp', 0),
                                    'structure_hash': cached_data.get('structure_hash', ''),
                                    'item_count': cached_data.get('item_count', 0)
                                })
                        except Exception:
                            continue
                except Exception as e:
                    logger.warning(f"Failed to export Direct LLM cache: {e}")
        
        all_patterns = patterns + direct_llm_patterns
        
        export_data = {
            'version': '1.0',
            'exported_at': time.time(),
            'tenant_id': tenant_id,
            'domain_filter': domain,
            'total_patterns': len(all_patterns),
            'patterns': all_patterns,
            'domains': sorted(list(set(p.get('domain') for p in all_patterns if p.get('domain'))))
        }
        
        json_str = json.dumps(export_data, indent=2, default=str)
        return Response(
            content=json_str,
            media_type="application/json",
            headers={
                "Content-Disposition": f"attachment; filename=paradocs-cache-export-{domain or 'all'}-{int(time.time())}.json"
            }
        )
    except Exception as e:
        logger.error(f"Cache export failed: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Cache export failed: {str(e)}"
        )


# ============================================================================
# TENANT PATTERN CACHE ENDPOINTS (Multi-tenant with public/private visibility)
# ============================================================================

class PatternVisibilityUpdate(BaseModel):
    visibility: str = Field(..., description="Visibility: 'public' or 'private'")

class PatternStoreRequest(BaseModel):
    domain: str
    fields: List[str]
    pattern_data: Dict[str, Any]
    visibility: str = "private"
    url: Optional[str] = None

@app.get("/api/v1/patterns/mine")
async def list_my_patterns(
    domain: Optional[str] = None,
    visibility: Optional[str] = None,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    List current user's cached patterns (private and public)
    """
    try:
        cache = get_tenant_pattern_cache()
        vis_filter = CacheVisibility(visibility) if visibility else None
        patterns = await cache.list_tenant_patterns(tenant_id, domain, vis_filter)
        
        return {
            "success": True,
            "patterns": patterns,
            "total": len(patterns),
            "tenant_id": tenant_id
        }
    except Exception as e:
        logger.error(f"Failed to list patterns: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/patterns/public")
async def list_public_patterns(
    domain: Optional[str] = None,
    limit: int = 100,
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    List all public patterns (shared by community)
    """
    try:
        cache = get_tenant_pattern_cache()
        patterns = await cache.list_public_patterns(domain, limit)
        
        return {
            "success": True,
            "patterns": patterns,
            "total": len(patterns)
        }
    except Exception as e:
        logger.error(f"Failed to list public patterns: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/patterns/store")
async def store_pattern(
    request: PatternStoreRequest,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Store a new extraction pattern
    """
    try:
        cache = get_tenant_pattern_cache()
        visibility = CacheVisibility(request.visibility)
        
        success = await cache.store_pattern(
            tenant_id=tenant_id,
            domain=request.domain,
            fields=request.fields,
            pattern_data=request.pattern_data,
            visibility=visibility,
            url=request.url
        )
        
        return {
            "success": success,
            "message": "Pattern stored successfully" if success else "Failed to store pattern"
        }
    except Exception as e:
        logger.error(f"Failed to store pattern: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/api/v1/patterns/visibility")
async def update_pattern_visibility(
    domain: str,
    fields: str,  # Comma-separated
    request: PatternVisibilityUpdate,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Update pattern visibility (make public or private)
    """
    try:
        cache = get_tenant_pattern_cache()
        fields_list = [f.strip() for f in fields.split(',')]
        visibility = CacheVisibility(request.visibility)
        
        success = await cache.update_visibility(tenant_id, domain, fields_list, visibility)
        
        return {
            "success": success,
            "message": f"Pattern visibility updated to {request.visibility}" if success else "Pattern not found"
        }
    except Exception as e:
        logger.error(f"Failed to update visibility: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/patterns/copy-public")
async def copy_public_pattern(
    domain: str,
    fields: str,  # Comma-separated
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Copy a public pattern to your private cache
    """
    try:
        cache = get_tenant_pattern_cache()
        fields_list = [f.strip() for f in fields.split(',')]
        
        success = await cache.copy_public_pattern(tenant_id, domain, fields_list)
        
        return {
            "success": success,
            "message": "Pattern copied to your cache" if success else "Public pattern not found"
        }
    except Exception as e:
        logger.error(f"Failed to copy pattern: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/v1/patterns")
async def delete_pattern(
    domain: str = Query(..., description="Domain of the pattern to delete"),
    fields: str = Query(..., description="Comma-separated list of fields"),
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Delete a pattern from your cache
    """
    try:
        logger.info(f"[{tenant_id}] Deleting pattern: domain={domain}, fields={fields}")
        cache = get_tenant_pattern_cache()
        fields_list = [f.strip() for f in fields.split(',')]
        
        logger.info(f"[{tenant_id}] Parsed fields: {fields_list}")
        
        success = await cache.delete_pattern(tenant_id, domain, fields_list)
        
        logger.info(f"[{tenant_id}] Delete result: {success}")
        
        return {
            "success": success,
            "message": "Pattern deleted" if success else "Pattern not found"
        }
    except Exception as e:
        logger.error(f"[{tenant_id}] Failed to delete pattern: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/patterns/stats")
async def get_pattern_stats(
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Get pattern cache statistics
    """
    try:
        cache = get_tenant_pattern_cache()
        stats = await cache.get_stats(tenant_id)
        
        return {
            "success": True,
            **stats
        }
    except Exception as e:
        logger.error(f"Failed to get stats: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# AGENTS (JOBS) ENDPOINTS
# ============================================================================

class CreateAgentRequest(BaseModel):
    type: str = Field(..., description="Agent type: 'web_scraping', 'document_processing', 'batch_scraping'")
    config: Dict[str, Any] = Field(..., description="Agent configuration")
    queue_immediately: bool = Field(default=True, description="Queue for execution immediately")

@app.post("/api/v1/agents")
async def create_agent(
    request: CreateAgentRequest,
    background_tasks: BackgroundTasks,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Create a new agent (async job)
    """
    try:
        manager = get_agent_manager()
        agent_type = AgentType(request.type)
        
        agent = await manager.create_agent(
            tenant_id=tenant_id,
            agent_type=agent_type,
            config=request.config,
            queue_immediately=request.queue_immediately
        )
        
        # If not using Cloud Tasks, execute in background
        if request.queue_immediately and not manager.cloud_tasks_enabled:
            background_tasks.add_task(execute_agent_task, agent.id, x_api_key)
        
        return {
            "success": True,
            "agent": agent.to_dict()
        }
    except Exception as e:
        logger.error(f"Failed to create agent: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/agents")
async def list_agents(
    status: Optional[str] = None,
    type: Optional[str] = None,
    limit: int = 50,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    List agents for current user
    """
    try:
        manager = get_agent_manager()
        status_filter = AgentStatus(status) if status else None
        type_filter = AgentType(type) if type else None
        
        agents = await manager.list_agents(tenant_id, status_filter, type_filter, limit)
        
        return {
            "success": True,
            "agents": [a.to_dict() for a in agents],
            "total": len(agents)
        }
    except Exception as e:
        logger.error(f"Failed to list agents: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

# NOTE: stats endpoint MUST come before {agent_id} to avoid "stats" being treated as an ID
@app.get("/api/v1/agents/stats")
async def get_agent_stats(
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Get agent statistics for current user
    """
    try:
        manager = get_agent_manager()
        stats = await manager.get_stats(tenant_id)
        
        return {
            "success": True,
            **stats
        }
    except Exception as e:
        logger.error(f"Failed to get agent stats: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/agents/{agent_id}")
async def get_agent(
    agent_id: str,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Get agent details and status
    """
    try:
        manager = get_agent_manager()
        agent = await manager.get_agent(agent_id)
        
        if not agent:
            raise HTTPException(status_code=404, detail="Agent not found")
        
        # Verify tenant ownership
        if agent.tenant_id != tenant_id:
            raise HTTPException(status_code=403, detail="Access denied")
        
        return {
            "success": True,
            "agent": agent.to_dict()
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get agent: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/agents/{agent_id}/cancel")
async def cancel_agent(
    agent_id: str,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Cancel a pending/queued agent
    """
    try:
        manager = get_agent_manager()
        agent = await manager.get_agent(agent_id)
        
        if not agent:
            raise HTTPException(status_code=404, detail="Agent not found")
        
        if agent.tenant_id != tenant_id:
            raise HTTPException(status_code=403, detail="Access denied")
        
        success = await manager.cancel_agent(agent_id)
        
        return {
            "success": success,
            "message": "Agent cancelled" if success else "Cannot cancel agent in current state"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to cancel agent: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/agents/{agent_id}/execute")
async def execute_agent_endpoint(
    agent_id: str,
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Execute agent (called by Cloud Tasks, Cloud Scheduler, or internal trigger)
    """
    try:
        await execute_agent_task(agent_id, x_api_key)
        return {"success": True, "message": "Agent execution started"}
    except Exception as e:
        logger.error(f"Failed to execute agent: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

class ScheduleAgentRequest(BaseModel):
    schedule: str = Field(..., description="Cron expression (e.g., '0 */6 * * *' for every 6 hours)")
    timezone: str = Field(default="UTC", description="Timezone for schedule")

@app.post("/api/v1/agents/{agent_id}/schedule")
async def schedule_agent(
    agent_id: str,
    request: ScheduleAgentRequest,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Schedule an agent to run periodically
    
    Common schedules:
    - "0 * * * *" - Every hour
    - "0 */6 * * *" - Every 6 hours
    - "0 0 * * *" - Daily at midnight
    - "0 0 * * 0" - Weekly on Sunday
    - "0 0 1 * *" - Monthly on the 1st
    """
    try:
        manager = get_agent_manager()
        agent = await manager.get_agent(agent_id)
        
        if not agent:
            raise HTTPException(status_code=404, detail="Agent not found")
        
        if agent.tenant_id != tenant_id:
            raise HTTPException(status_code=403, detail="Access denied")
        
        success = await manager.schedule_agent(agent_id, request.schedule, request.timezone)
        
        return {
            "success": success,
            "message": f"Agent scheduled: {request.schedule}" if success else "Failed to schedule agent (Cloud Scheduler may not be available)"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to schedule agent: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/v1/agents/{agent_id}/schedule")
async def unschedule_agent(
    agent_id: str,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Remove schedule from an agent
    """
    try:
        manager = get_agent_manager()
        agent = await manager.get_agent(agent_id)
        
        if not agent:
            raise HTTPException(status_code=404, detail="Agent not found")
        
        if agent.tenant_id != tenant_id:
            raise HTTPException(status_code=403, detail="Access denied")
        
        success = await manager.unschedule_agent(agent_id)
        
        return {
            "success": success,
            "message": "Schedule removed" if success else "Agent was not scheduled"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to unschedule agent: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

class CreateAgentFromCacheRequest(BaseModel):
    domain: str = Field(..., description="Domain from cache")
    fields: List[str] = Field(..., description="Fields to extract")
    url: str = Field(..., description="URL to scrape")
    visibility: str = Field(default="private", description="Cache visibility")
    schedule: Optional[str] = Field(default=None, description="Optional cron schedule")

@app.post("/api/v1/agents/from-cache")
async def create_agent_from_cache(
    request: CreateAgentFromCacheRequest,
    background_tasks: BackgroundTasks,
    tenant_id: str = Depends(get_tenant_id),
    x_api_key: Optional[str] = Header(None, alias="X-API-Key")
):
    """
    Create an agent from a cached pattern
    """
    try:
        manager = get_agent_manager()
        
        agent = await manager.create_from_cache(
            tenant_id=tenant_id,
            domain=request.domain,
            fields=request.fields,
            url=request.url,
            visibility=request.visibility,
            schedule=request.schedule
        )
        
        # If no schedule and not using Cloud Tasks, execute in background
        if not request.schedule and not manager.cloud_tasks_enabled:
            background_tasks.add_task(execute_agent_task, agent.id, x_api_key)
        
        return {
            "success": True,
            "agent": agent.to_dict(),
            "message": f"Agent created from cache ({request.domain})" + (f" with schedule: {request.schedule}" if request.schedule else "")
        }
    except Exception as e:
        logger.error(f"Failed to create agent from cache: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

async def execute_agent_task(agent_id: str, api_key: Optional[str] = None):
    """
    Background task to execute an agent
    """
    manager = get_agent_manager()
    agent = await manager.get_agent(agent_id)
    
    if not agent:
        logger.error(f"Agent not found: {agent_id}")
        return
    
    try:
        # Update status to running
        await manager.update_progress(agent_id, 10, "Starting execution...")
        
        config = agent.config
        
        if agent.type == AgentType.WEB_SCRAPING:
            # Execute web scraping
            await manager.update_progress(agent_id, 20, "Fetching page...")
            
            key = api_key or os.getenv("OPENAI_API_KEY") or os.getenv("GEMINI_API_KEY")
            scraper = get_scraper(key, config.get("mode", "hybrid"), redis_cache=get_redis_cache())
            
            await manager.update_progress(agent_id, 40, "Extracting data...")
            
            result = await scraper.scrape(
                url=config.get("url"),
                fields=config.get("fields", []),
                scroll_to_bottom=config.get("scroll_to_bottom", False),
                wait_for_selector=config.get("wait_for_selector")
            )
            
            await manager.update_progress(agent_id, 80, "Processing results...")
            
            # Store pattern if extraction was successful
            if result.get("data"):
                pattern_cache = get_tenant_pattern_cache()
                from urllib.parse import urlparse
                domain = urlparse(config.get("url")).netloc
                
                await pattern_cache.store_pattern(
                    tenant_id=agent.tenant_id,
                    domain=domain,
                    fields=config.get("fields", []),
                    pattern_data={"source": result.get("source"), "item_count": len(result.get("data", []))},
                    visibility=CacheVisibility.PRIVATE,
                    url=config.get("url")
                )
            
            await manager.complete_agent(agent_id, result)
            
        elif agent.type == AgentType.DOCUMENT_PROCESSING:
            # Execute document processing
            await manager.update_progress(agent_id, 20, "Processing document...")
            
            # Document processing logic would go here
            # For now, mark as completed with placeholder
            await manager.complete_agent(agent_id, {"message": "Document processing not yet implemented in agent mode"})
            
        elif agent.type == AgentType.BATCH_SCRAPING:
            # Execute batch scraping
            urls = config.get("urls", [])
            fields = config.get("fields", [])
            results = []
            
            key = api_key or os.getenv("OPENAI_API_KEY") or os.getenv("GEMINI_API_KEY")
            scraper = get_scraper(key, config.get("mode", "hybrid"), redis_cache=get_redis_cache())
            
            for i, url in enumerate(urls):
                progress = int(20 + (i / len(urls)) * 60)
                await manager.update_progress(agent_id, progress, f"Scraping {i+1}/{len(urls)}: {url[:50]}...")
                
                try:
                    result = await scraper.scrape(url=url, fields=fields)
                    results.append({"url": url, "data": result.get("data", []), "success": True})
                except Exception as e:
                    results.append({"url": url, "error": str(e), "success": False})
            
            await manager.complete_agent(agent_id, {"results": results, "total": len(results)})
        
        else:
            await manager.fail_agent(agent_id, f"Unknown agent type: {agent.type}")
            
    except Exception as e:
        logger.error(f"Agent execution failed: {e}", exc_info=True)
        await manager.fail_agent(agent_id, str(e))


# ============================================================================
# BROWSER SESSION ENDPOINTS (WebSocket-based live preview)
# ============================================================================

from fastapi import WebSocket, WebSocketDisconnect
from api.browser_session import get_session_manager, shutdown_session_manager, BrowserSession

class BrowserSessionRequest(BaseModel):
    proxy_config: Optional[Dict[str, Any]] = None
    headless: bool = True
    viewport: Optional[Dict[str, int]] = None

class NavigateRequest(BaseModel):
    url: str
    wait_for: str = "domcontentloaded"
    timeout: int = 60000

class ClickRequest(BaseModel):
    selector: str
    button: str = "left"

class ScrollRequest(BaseModel):
    direction: str = "down"
    amount: int = 500

class SelectElementRequest(BaseModel):
    selector: str
    field_name: str

@app.post("/api/v1/browser/session")
async def create_browser_session(
    request: BrowserSessionRequest,
    tenant_id: str = Depends(get_tenant_id)
):
    """
    Create a new browser session for live preview.
    Returns session_id for subsequent operations.
    """
    try:
        manager = await get_session_manager()
        
        # Convert proxy config if needed
        proxy_config = convert_proxy_config(request.proxy_config) if request.proxy_config else None
        
        session = await manager.create_session(
            tenant_id=tenant_id,
            proxy_config=proxy_config,
            headless=request.headless,
            viewport=request.viewport
        )
        
        return {
            "success": True,
            "session_id": session.id,
            "message": "Browser session created"
        }
    except Exception as e:
        logger.error(f"Failed to create browser session: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/v1/browser/session/{session_id}")
async def close_browser_session(
    session_id: str,
    tenant_id: str = Depends(get_tenant_id)
):
    """Close a browser session"""
    try:
        manager = await get_session_manager()
        success = await manager.close_session(session_id)
        
        return {
            "success": success,
            "message": "Session closed" if success else "Session not found"
        }
    except Exception as e:
        logger.error(f"Failed to close browser session: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/browser/session/{session_id}/navigate")
async def browser_navigate(
    session_id: str,
    request: NavigateRequest,
    tenant_id: str = Depends(get_tenant_id)
):
    """Navigate to a URL in the browser session"""
    try:
        manager = await get_session_manager()
        result = await manager.navigate(
            session_id=session_id,
            url=request.url,
            wait_for=request.wait_for,
            timeout=request.timeout
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Navigation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/browser/session/{session_id}/click")
async def browser_click(
    session_id: str,
    request: ClickRequest,
    tenant_id: str = Depends(get_tenant_id)
):
    """Click an element in the browser session"""
    try:
        manager = await get_session_manager()
        result = await manager.click(
            session_id=session_id,
            selector=request.selector,
            button=request.button
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Click failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/browser/session/{session_id}/scroll")
async def browser_scroll(
    session_id: str,
    request: ScrollRequest,
    tenant_id: str = Depends(get_tenant_id)
):
    """Scroll the page in the browser session"""
    try:
        manager = await get_session_manager()
        result = await manager.scroll(
            session_id=session_id,
            direction=request.direction,
            amount=request.amount
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Scroll failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/browser/session/{session_id}/screenshot")
async def browser_screenshot(
    session_id: str,
    tenant_id: str = Depends(get_tenant_id)
):
    """Get current screenshot of the browser session"""
    try:
        manager = await get_session_manager()
        screenshot = await manager.get_screenshot(session_id)
        
        if screenshot:
            return {"success": True, "screenshot": screenshot}
        else:
            raise HTTPException(status_code=404, detail="Session not found or inactive")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Screenshot failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/browser/session/{session_id}/html")
async def browser_html(
    session_id: str,
    tenant_id: str = Depends(get_tenant_id)
):
    """Get current HTML of the browser session"""
    try:
        manager = await get_session_manager()
        html = await manager.get_html(session_id)
        
        if html:
            return {"success": True, "html": html, "size": len(html)}
        else:
            raise HTTPException(status_code=404, detail="Session not found or inactive")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"HTML fetch failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/browser/session/{session_id}/select")
async def browser_select_element(
    session_id: str,
    request: SelectElementRequest,
    tenant_id: str = Depends(get_tenant_id)
):
    """Select an element as a field for extraction"""
    try:
        manager = await get_session_manager()
        result = await manager.select_element(
            session_id=session_id,
            selector=request.selector,
            field_name=request.field_name
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Element selection failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/browser/session/{session_id}/selections")
async def browser_get_selections(
    session_id: str,
    tenant_id: str = Depends(get_tenant_id)
):
    """Get all selected elements for the session"""
    try:
        manager = await get_session_manager()
        selections = await manager.get_selected_elements(session_id)
        return {"success": True, "selections": selections}
    except Exception as e:
        logger.error(f"Get selections failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/v1/browser/session/{session_id}/selections")
async def browser_clear_selections(
    session_id: str,
    tenant_id: str = Depends(get_tenant_id)
):
    """Clear all selected elements"""
    try:
        manager = await get_session_manager()
        success = await manager.clear_selections(session_id)
        return {"success": success}
    except Exception as e:
        logger.error(f"Clear selections failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/browser/session/{session_id}/evaluate")
async def browser_evaluate(
    session_id: str,
    script: str = Form(...),
    tenant_id: str = Depends(get_tenant_id)
):
    """Execute JavaScript in the browser session"""
    try:
        manager = await get_session_manager()
        result = await manager.evaluate(session_id, script)
        return {"success": True, "result": result}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Script evaluation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# App lifecycle events
@app.on_event("shutdown")
async def app_shutdown():
    """Cleanup on app shutdown"""
    await shutdown_session_manager()


@app.post("/api/v1/save-pattern")
async def save_pattern_endpoint(
    request: SavePatternRequest,
    tenant_id: str = Depends(get_tenant_id),
    pattern_cache: TenantPatternCache = Depends(get_tenant_pattern_cache)
):
    """Save an extraction pattern for a tenant"""
    from urllib.parse import urlparse
    domain = urlparse(request.url).netloc
    
    logger.info(f"[{tenant_id}] Saving pattern for domain: {domain}")
    
    visibility = CacheVisibility.PUBLIC if request.visibility == "public" else CacheVisibility.PRIVATE
    
    success = await pattern_cache.store_pattern(
        tenant_id=tenant_id,
        domain=domain,
        fields=request.fields,
        pattern_data=request.pattern_data,
        visibility=visibility,
        url=request.url
    )
    
    logger.info(f"[{tenant_id}] Pattern storage result: {success}")
    
    if not success:
        raise HTTPException(status_code=500, detail="Failed to store pattern")
        
    return {"success": True, "message": "Pattern saved successfully"}

@app.get("/api/v1/list-patterns")
async def list_patterns_endpoint(
    domain: Optional[str] = None,
    tenant_id: str = Depends(get_tenant_id),
    pattern_cache: TenantPatternCache = Depends(get_tenant_pattern_cache)
):
    """List all patterns for a tenant"""
    logger.info(f"[{tenant_id}] Listing patterns for domain: {domain}")
    patterns = await pattern_cache.list_tenant_patterns(tenant_id, domain=domain)
    logger.info(f"[{tenant_id}] Found {len(patterns)} patterns")
    return {"success": True, "patterns": patterns}

@app.post("/api/v1/delete-pattern")
async def delete_pattern_endpoint(
    request: DeletePatternRequest,
    tenant_id: str = Depends(get_tenant_id),
    pattern_cache: TenantPatternCache = Depends(get_tenant_pattern_cache)
):
    """Delete an extraction pattern"""
    success = await pattern_cache.delete_pattern(
        tenant_id=tenant_id,
        domain=request.domain,
        fields=request.fields
    )
    
    if not success:
        raise HTTPException(status_code=404, detail="Pattern not found or could not be deleted")
        
    return {"success": True, "message": "Pattern deleted successfully"}


@app.post("/api/v1/generate-python-code")
async def generate_python_code_endpoint(
    request: GenerateCodeRequest,
    tenant_id: str = Depends(get_tenant_id)
):
    """Generate a standalone Python script for the extraction pattern"""
    try:
        from universal_scraper.core.code_generator import PythonCodeGenerator
        
        generator = PythonCodeGenerator()
        code = generator.generate_script(
            url=request.url,
            fields=request.fields,
            selectors=request.selectors,
            target=request.target
        )
        
        from urllib.parse import urlparse
        return {
            "success": True,
            "code": code,
            "filename": f"scraper_{urlparse(request.url).netloc.replace('.', '_')}.py"
        }
    except Exception as e:
        logger.error(f"Code generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Code generation failed: {str(e)}")


if __name__ == "__main__":
    port = int(os.getenv("PORT", "8080"))
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=port,
        log_level="info",
        access_log=True
    )
